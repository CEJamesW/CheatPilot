from __future__ import annotations

import re
from zipfile import ZIP_DEFLATED, ZipFile
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r'C:\Users\Administrator\Desktop\CheatPilot')
TEMPLATE = ROOT / '信息技术实践与拓展-报告模版.docx'
OUT = ROOT / '信息技术实践与拓展-最终报告-22组-CheatPilot.docx'
IMG_DIR = ROOT / 'report_assets'
IMG_DIR.mkdir(exist_ok=True)

TITLE = 'CheatPilot：由 LLM 驱动的实时内存修改系统'
TEAM_NO = '22'
LEADER = '王斯民'
MEMBERS = ['孙小然', '杜雨金', '宋震浩', '贾奇东']
GITHUB_URL = 'https://github.com/CEJamesW/CheatPilot'


def font_path() -> str | None:
    candidates = [
        r'C:\Windows\Fonts\msyh.ttc',
        r'C:\Windows\Fonts\simhei.ttf',
        r'C:\Windows\Fonts\simsun.ttc',
    ]
    for p in candidates:
        if Path(p).exists():
            return p
    return None

FONT = font_path()


def get_font(size: int, bold: bool = False):
    if FONT:
        return ImageFont.truetype(FONT, size=size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill=(17,24,39)):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2-x1-w)/2, y1 + (y2-y1-h)/2), text, font=font, fill=fill, align='center', spacing=6)


def rounded_rect(draw, xy, radius, fill, outline=None, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def make_module_image(path: Path):
    img = Image.new('RGB', (1400, 780), 'white')
    d = ImageDraw.Draw(img)
    title_f = get_font(42)
    box_f = get_font(28)
    small_f = get_font(22)
    d.text((55, 42), 'CheatPilot 功能模块图', font=title_f, fill=(17,24,39))
    boxes = [
        (80, 145, 380, 285, '自然语言入口\nCLI / API / UI', (219,234,254)),
        (550, 145, 850, 285, 'LLM Agent\n工具规划 / 思考', (220,252,231)),
        (1020, 145, 1320, 285, '工具执行层\nComposite Executor', (254,243,199)),
        (180, 440, 480, 580, 'Cheat Engine MCP\n扫描 / 写入 / 读回', (254,226,226)),
        (560, 440, 860, 580, '会话状态管理\n候选地址 / 待写入', (237,233,254)),
        (940, 440, 1240, 580, '本地工具\n进程 / 文件 / 命令', (224,242,254)),
    ]
    for x1,y1,x2,y2,text,color in boxes:
        rounded_rect(d, (x1,y1,x2,y2), 22, color, outline=(100,116,139), width=3)
        draw_centered(d, (x1,y1,x2,y2), text, box_f)
    arrows = [((380,215),(550,215)), ((850,215),(1020,215)), ((1170,285),(330,440)), ((700,285),(710,440)), ((1170,285),(1090,440))]
    for a,b in arrows:
        d.line((a,b), fill=(71,85,105), width=5)
        ax, ay = a; bx, by = b
        # simple arrow head
        d.polygon([(bx,by), (bx-18,by-10), (bx-18,by+10)] if abs(by-ay)<20 and bx>ax else [(bx,by), (bx-10,by-18), (bx+10,by-18)], fill=(71,85,105))
    d.text((86, 680), '说明：用户的全部输入首先进入 LLM Agent，由模型根据上下文选择真实工具调用；内存操作统一经 Cheat Engine MCP 执行。', font=small_f, fill=(71,85,105))
    img.save(path)


def make_flow_image(path: Path):
    img = Image.new('RGB', (1400, 850), 'white')
    d = ImageDraw.Draw(img)
    title_f = get_font(42)
    f = get_font(25)
    small_f = get_font(20)
    d.text((55, 38), '系统流程图', font=title_f, fill=(17,24,39))
    steps = [
        ('用户输入任务', '目标进程、当前值、目标值'),
        ('LLM 判断缺失信息', '需要时询问当前可见值'),
        ('附加目标进程', '使用 list_processes / attach_process'),
        ('首次精确扫描', 'scan_exact_value'),
        ('候选是否唯一', '根据 CE total 判断'),
        ('多轮 next scan', '引导用户改变同一数值'),
        ('写入并读回确认', 'write_value + read_value'),
        ('输出结果', '地址 / 基址 / 下一步'),
    ]
    x = 120; y = 125; w = 360; h = 92; gap_y = 38
    positions = []
    for i, (name, desc) in enumerate(steps):
        col = 0 if i < 4 else 1
        row = i if i < 4 else i - 4
        px = x + col * 680
        py = y + row * (h + gap_y)
        positions.append((px, py, px+w, py+h))
        rounded_rect(d, (px,py,px+w,py+h), 18, (240,249,255) if i % 2 == 0 else (245,253,244), outline=(100,116,139), width=3)
        d.text((px+24, py+16), name, font=f, fill=(17,24,39))
        d.text((px+24, py+54), desc, font=small_f, fill=(71,85,105))
    def connect(i,j):
        x1,y1,x2,y2 = positions[i]
        a=(x1+w//2,y2); b=(positions[j][0]+w//2, positions[j][1])
        if j == 4:
            a=(x2,y1+h//2); b=(positions[j][0], positions[j][1]+h//2)
        d.line((a,b), fill=(37,99,235), width=4)
    for i in range(3): connect(i,i+1)
    connect(3,4)
    for i in range(4,7): connect(i,i+1)
    # loop arrow from uniqueness to next scan and back concept
    d.arc((860, 380, 1260, 650), start=20, end=330, fill=(220,38,38), width=4)
    d.text((885, 650), '候选不唯一时继续变化数值并筛选', font=small_f, fill=(185,28,28))
    d.text((80, 770), '关键控制点：不根据预览地址盲写；只有 Cheat Engine 返回唯一总匹配且读回确认后，才向用户报告写入成功。', font=small_f, fill=(71,85,105))
    img.save(path)


def make_ui_image(path: Path):
    img = Image.new('RGB', (1400, 820), (248,250,252))
    d = ImageDraw.Draw(img)
    title_f = get_font(42)
    f = get_font(25)
    mono = get_font(22)
    d.text((55, 40), '桌面聊天界面原型', font=title_f, fill=(17,24,39))
    rounded_rect(d, (220, 120, 1180, 705), 16, 'white', outline=(148,163,184), width=3)
    d.rectangle((220,120,1180,185), fill=(15,23,42))
    d.text((250, 138), 'CheatPilot', font=f, fill='white')
    bubbles = [
        ('user> 附加到 game.exe', 260, 225, (219,234,254)),
        ('assistant> 已附加到目标进程。', 260, 295, (220,252,231)),
        ('user> 当前金币是150，改成9999并打印地址', 260, 365, (219,234,254)),
        ('assistant> 扫描到多个候选。请让金币变化一次后告诉我新值。', 260, 435, (220,252,231)),
        ('status> 思考中...', 260, 505, (254,249,195)),
    ]
    for text,x,y,color in bubbles:
        rounded_rect(d, (x,y,1140,y+48), 12, color, outline=(203,213,225), width=2)
        d.text((x+18,y+12), text, font=mono, fill=(15,23,42))
    rounded_rect(d, (250, 625, 1030, 675), 12, (241,245,249), outline=(203,213,225), width=2)
    d.text((270, 638), '输入自然语言任务...', font=mono, fill=(100,116,139))
    rounded_rect(d, (1048, 625, 1148, 675), 12, (37,99,235), outline=(37,99,235), width=2)
    d.text((1078, 638), '发送', font=mono, fill='white')
    img.save(path)


def make_effect_image(path: Path):
    img = Image.new('RGB', (1400, 760), 'white')
    d = ImageDraw.Draw(img)
    title_f = get_font(42)
    f = get_font(24)
    mono = get_font(21)
    d.text((55, 38), '实现效果示意图', font=title_f, fill=(17,24,39))
    cols = [(60,130,430,610,'CLI 交互'), (515,130,885,610,'API 服务'), (970,130,1340,610,'CE MCP 后端')]
    for x1,y1,x2,y2,title in cols:
        rounded_rect(d, (x1,y1,x2,y2), 16, (248,250,252), outline=(148,163,184), width=3)
        d.text((x1+24,y1+22), title, font=f, fill=(17,24,39))
    d.text((85,195), 'python -m cheatpilot -i\n\nuser> 当前金币是150\nassistant> 扫描候选...\nuser> 现在金币是120\nassistant> 写入9999并读回', font=mono, fill=(30,41,59), spacing=8)
    d.text((540,195), 'POST /chat\n{\n  "session_id": "demo",\n  "message": "现在金币是120"\n}\n\n返回 assistant_message\n和结构化 results', font=mono, fill=(30,41,59), spacing=8)
    d.text((995,195), 'openProcess(pid)\nscan_all(value,type)\nnext_scan(value)\nwrite_integer(addr,value)\nread_integer(addr)\n\n真实工具结果回传给 LLM', font=mono, fill=(30,41,59), spacing=8)
    # arrows
    d.line((430,370,515,370), fill=(37,99,235), width=5)
    d.polygon([(515,370),(495,360),(495,380)], fill=(37,99,235))
    d.line((885,370,970,370), fill=(37,99,235), width=5)
    d.polygon([(970,370),(950,360),(950,380)], fill=(37,99,235))
    d.text((78,690), '实现入口覆盖 CLI、API 和桌面 UI；所有内存读写统一经 Cheat Engine MCP，结果再由 LLM 汇总为中文回复。', font=f, fill=(71,85,105))
    img.save(path)


def set_cell_text(cell, text: str):
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            r.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name='宋体', size=10.5, bold=False, color=None):
    if name is not None:
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_paragraph(p):
    for r in list(p.runs):
        p._element.remove(r._element)


def set_para_text(p, text, *, size=10.5, bold=False, align=None, font='宋体'):
    clear_paragraph(p)
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold)
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = None
    return p


def add_para(doc, text='', style=None, *, size=10.5, bold=False, align=None, first_line=True, before=0, after=0, line=1.25):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.add_run(text)
    if align is not None:
        p.alignment = align
    if first_line and text:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0,0)
    cell.text = ''
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F3F4F6')
    cell._tc.get_or_add_tcPr().append(shading)
    p = cell.paragraphs[0]
    for line in code.strip('\n').split('\n'):
        run = p.add_run(line)
        set_run_font(run, name='Consolas', size=8.5)
        p.add_run('\n')
    return table


def add_picture_with_caption(doc, image_path: Path, caption: str, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(0)
        p.add_run(f'• {item}')
        p.paragraph_format.space_after = Pt(2)


def fill_underlined_runs(p, text: str):
    underline_runs = [r for r in p.runs if r.font.underline]
    if not underline_runs:
        p.add_run(text)
        return
    underline_runs[0].text = text
    for run in underline_runs[1:]:
        run.text = ''


def replace_paragraph_text_preserve_first_run(p, text: str):
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for run in p.runs[1:]:
        run.text = ''


def fill_group_number(p, team_no: str):
    for idx, run in enumerate(p.runs):
        if run.text == '第':
            for candidate in p.runs[idx + 1:]:
                if candidate.font.underline and '组' not in candidate.text:
                    candidate.text = f'  {team_no}  '
                    return
    fill_underlined_runs(p, f'第  {team_no}  组')


def fill_date(p, date_text: str):
    if len(p.runs) >= len(date_text):
        for run, char in zip(p.runs, date_text):
            run.text = char
        for run in p.runs[len(date_text):]:
            run.text = ''
        return
    replace_paragraph_text_preserve_first_run(p, date_text)


def fill_cover(doc: Document):
    title_para = doc.paragraphs[9]
    fill_underlined_runs(title_para, TITLE)
    if len(title_para.runs) > 3:
        title_para.runs[2].text = ' '
        title_para.runs[3].text = ''
    fill_group_number(doc.paragraphs[14], TEAM_NO)
    fill_underlined_runs(doc.paragraphs[15], LEADER)
    for idx, member in zip(range(16, 20), MEMBERS):
        fill_underlined_runs(doc.paragraphs[idx], member)
    fill_date(doc.paragraphs[24], '2026年6月')


def remove_body_after_material_link(doc: Document):
    # keep cover, scoring rules and material link heading through paragraph 29; remove old placeholder paragraphs after it.
    for idx in range(len(doc.paragraphs)-1, 29, -1):
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)
    # remove sample test table only, keep scoring table
    while len(doc.tables) > 1:
        t = doc.tables[-1]
        t._element.getparent().remove(t._element)


def _xml_local_name(element) -> str:
    tag = getattr(element, 'tag', '')
    return tag.rsplit('}', 1)[-1] if isinstance(tag, str) else ''


def _shape_run_for_template_note(text_element):
    current = text_element
    shape_seen = False
    fallback = None
    while current is not None:
        if _xml_local_name(current) in {'AlternateContent', 'drawing', 'pict'}:
            shape_seen = True
            fallback = current
        parent = current.getparent()
        if shape_seen and parent is not None and _xml_local_name(current) == 'r' and _xml_local_name(parent) == 'p':
            return current
        current = parent
    return fallback


def remove_template_note_shapes(docx_path: Path):
    hints = ('最多5人', '提示删除', '第1位同学')
    tmp_path = docx_path.with_suffix(docx_path.suffix + '.tmp')
    with ZipFile(docx_path, 'r') as src:
        document_xml = src.read('word/document.xml')
        root = etree.fromstring(document_xml)
        victims = []
        for element in root.iter():
            if _xml_local_name(element) != 't' or not element.text:
                continue
            if any(hint in element.text for hint in hints):
                victim = _shape_run_for_template_note(element)
                if victim is not None and not any(victim is item for item in victims):
                    victims.append(victim)
        for victim in victims:
            parent = victim.getparent()
            if parent is not None:
                parent.remove(victim)
        cleaned_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
        with ZipFile(tmp_path, 'w', ZIP_DEFLATED) as dst:
            for item in src.infolist():
                data = cleaned_xml if item.filename == 'word/document.xml' else src.read(item.filename)
                dst.writestr(item, data)
    tmp_path.replace(docx_path)


def add_test_table(doc: Document):
    add_para(doc, '表2.1 系统测试用例', align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    rows = [
        ['1', 'LLM tool calling', '运行 check_llm_tooluse，要求模型调用 session_status', '自然语言工具调用请求', '返回 OpenAI-compatible tool_calls', '通过'],
        ['2', '进程附加校验', '使用错误或过期进程名执行附加', 'game.exe / PID', '附加失败时停止后续扫描，避免误写进程', '通过'],
        ['3', '首次精确扫描', '输入当前可见数值并调用 scan_exact_value', '金币=150', '保存候选地址、总数和标签状态', '通过'],
        ['4', '多轮 next scan', '候选过多时改变数值后继续筛选', '金币=120', '沿用同一标签继续缩小候选集合', '通过'],
        ['5', '唯一地址写入', '候选唯一后执行写入并读回', '目标值=9999', '写入成功且读回值一致后报告成功', '通过'],
        ['6', 'API 多会话隔离', '两个 session 同时请求 CE 操作', 'session_id=s1/s2', 'CE 后端占用被正确识别，普通聊天不被阻塞', '通过'],
        ['7', '异常响应处理', '构造 malformed LLM 响应或 429/503', '异常 HTTP/JSON 响应', '返回明确中文错误提示，不静默失败', '通过'],
    ]
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['用例编号', '测试项', '测试步骤', '输入信号', '预期结果', '测试结果']
    for i,h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h)
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)


def build_report():
    module_img = IMG_DIR / 'figure_2_1_modules.png'
    ui_img = IMG_DIR / 'figure_2_2_ui.png'
    flow_img = IMG_DIR / 'figure_2_3_flow.png'
    effect_img = IMG_DIR / 'figure_2_5_effect.png'
    make_module_image(module_img)
    make_ui_image(ui_img)
    make_flow_image(flow_img)
    make_effect_image(effect_img)

    doc = Document(TEMPLATE)
    fill_cover(doc)
    if len(doc.paragraphs) > 29:
        replace_paragraph_text_preserve_first_run(doc.paragraphs[29], f'相关材料（电子报告及项目源码压缩包）的 GitHub 链接：{GITHUB_URL}')
    remove_body_after_material_link(doc)

    add_heading(doc, '1 技术调研报告', 1)
    add_heading(doc, '1.1 学习总结', 2)
    add_heading(doc, '1.1.1 内容简介', 3)
    add_para(doc, '本次实践围绕“由 LLM 驱动的实时内存修改系统”展开，项目名称为 CheatPilot。项目的基本目标不是编写一个固定流程脚本，而是探索如何让大语言模型在自然语言对话中理解用户目标，并通过标准化工具调用完成进程附加、内存扫描、候选筛选、数值写入、读回确认和地址输出等连续操作。传统内存修改工具需要使用者熟悉扫描类型、数值类型、候选地址缩小、指针和模块地址等概念；CheatPilot 试图把这些操作封装为可被 LLM 调度的工具，让使用者用“当前某个数值是多少，希望改成多少”这样的自然语言描述完成任务。')
    add_para(doc, '在技术调研阶段，小组重点学习了 OpenAI-compatible Chat Completions、tool calling、MCP 协议、Cheat Engine Lua Bridge、Windows 进程枚举、FastAPI 服务、桌面 UI 和 Python 文档化测试等内容。项目采用 Python 作为主要开发语言，核心由 ToolUseChatAgent、CompositeExecutor、CheatEngineMCPExecutor 和 LocalToolExecutor 组成。LLM 负责观察上下文并决定下一步工具调用；CompositeExecutor 根据 action 类型把本地文件、命令、进程查询等工具路由到本地执行器，把所有内存读写统一路由到 Cheat Engine MCP 执行器。这样的分层使 Agent 既能读取项目文件、列出进程、运行检查脚本，又能把真正的内存操作交给 Cheat Engine MCP 后端。')
    add_para(doc, '通过调研我们认识到，构建一个可用的 AI 工具并不只是调用模型接口。首先，模型输出必须被约束为可验证的工具调用，不能直接相信自然语言结论；其次，工具结果需要再次反馈给模型，让模型基于真实观察继续推理，而不是依赖本地硬编码流程；再次，外部服务可能出现限流、模型不可用、响应格式不兼容等问题，因此需要清晰的异常提示和重试机制。项目中曾遇到模型返回 429、503、无 tool_calls、旧式 function_call、JSON 格式异常等问题，最终通过重试、兼容解析、错误格式化和回归测试逐步提升稳定性。')
    add_para(doc, '本项目还让小组理解了状态管理的重要性。一次内存扫描往往不能直接得到唯一地址，Agent 必须保存标签、候选地址、总匹配数、上一次数值、待写入目标值和待打印地址等状态。当用户下一轮只说“现在是 120”时，系统要能根据唯一活跃会话推断这是同一个数值的 next scan，而不是重新开始。为避免误写，系统只在 Cheat Engine 明确返回唯一总匹配时写入，并在写入后读回确认。这个设计让我们体会到，AI Agent 的“智能”不仅体现在语言理解上，也体现在工具边界、状态连续性和执行确认机制上。')
    add_para(doc, '从工程角度看，CheatPilot 同时涉及 CLI、API 和桌面 UI 三种入口。CLI 适合快速调试，API 适合未来接入 Web 或其他 Agent 系统，桌面 UI 则提供直接对话窗口。项目通过统一的 AgentResponse 和 formatter 保证不同入口的返回信息一致。测试方面，项目包含 tool-agent、API 会话、扫描计数、数值类型、错误格式化等多类回归测试，避免在后续修改中破坏核心流程。总体而言，本次实践让小组从单点功能实现扩展到系统级设计，学习了 LLM 工具调用、真实后端集成、错误恢复和项目文档化等综合能力。')

    add_heading(doc, '1.1.2 难点和解决办法', 3)
    add_bullets(doc, [
        '难点一：LLM 输出不稳定。解决办法是定义严格 tool schema，兼容标准 tool_calls 和旧式 function_call，并在工具参数解析失败时把错误作为工具观察反馈给模型。',
        '难点二：内存扫描候选不唯一。解决办法是保存多轮扫描状态，让 Agent 引导用户改变同一数值后继续 next_scan，直到 Cheat Engine 返回唯一总匹配。',
        '难点三：进程附加可能误连旧进程。解决办法是在 attach 后读取进程信息并校验进程名或 PID，不匹配时停止后续扫描和写入。',
        '难点四：模型服务可能限流或不可用。解决办法是增加请求超时、自动重试和用户可读错误提示，并用检查脚本验证模型是否支持 tool calling。',
        '难点五：桌面、API、CLI 多入口状态一致。解决办法是统一 Agent 和 Executor 层，API 使用 session_id 管理不同会话，并对 CE 后端单实例占用进行保护。',
    ])

    add_heading(doc, '1.1.3 学习案例', 3)
    add_para(doc, '以“当前金币是 150，帮我改成 9999，并打印地址”为例，Agent 首先根据目标程序名称决定是否需要列出进程并附加目标进程；随后调用 scan_exact_value 扫描当前值。若返回候选过多，Agent 不会直接写入，而是提示用户让金币数值变化并报告新值。用户报告“现在金币是 120”后，Agent 调用 next_scan 继续筛选。只有当候选地址唯一且总匹配数为 1 时，系统才调用 write_value 写入 9999，并使用 read_value 或内部读回逻辑确认结果。最终回复中会给出写入结果、地址或基址信息。这一案例体现了 LLM 决策、工具调用、状态保存和真实执行结果确认的完整闭环。')

    add_heading(doc, '2 项目开发报告', 1)
    add_heading(doc, '2.1 项目简介', 2)
    add_para(doc, f'CheatPilot 是第 {TEAM_NO} 组完成的自然语言内存修改 Agent。系统以 LLM tool calling 为核心，以 Cheat Engine MCP 为真实执行后端，支持用户通过 CLI、API 或桌面 UI 输入自然语言任务。项目公开仓库地址为 {GITHUB_URL}。')
    add_para(doc, '项目的主要创新点在于把传统需要手动操作的扫描流程转化为对话式 Agent 流程。系统并不把“先扫、再缩、唯一后写”写死为简单脚本，而是把扫描、写入、读回、列进程、读取文件、执行命令等能力暴露为工具，由 LLM 根据上下文选择合适工具。Executor 层负责保护真实执行边界，保证内存操作均通过 Cheat Engine MCP 执行。')

    add_heading(doc, '2.2 需求分析', 2)
    add_heading(doc, '2.2.1 功能模块', 3)
    add_picture_with_caption(doc, module_img, '图2.1 功能模块图')
    add_heading(doc, '2.2.2 需求说明', 3)
    add_bullets(doc, [
        '自然语言输入需求：系统需要接受中文或英文任务描述，并把所有用户输入优先交给 LLM 处理。',
        '进程定位需求：当用户给出窗口名、程序名或 PID 时，系统需要能够列出候选进程并附加正确目标。',
        '数值扫描需求：系统需要支持整数和浮点数值扫描，保存标签和候选地址状态。',
        '多轮筛选需求：候选不唯一时，系统需要提示用户改变数值并继续 next scan。',
        '写入确认需求：只有在唯一候选地址下写入，并通过读回确认后，系统才报告成功。',
        '多入口需求：系统需要提供 CLI、API 和桌面 UI，便于不同使用方式接入。',
        '异常处理需求：模型限流、响应异常、MCP 超时、进程不匹配等情况必须给出明确提示。',
    ])

    add_heading(doc, '2.3 系统设计过程', 2)
    add_heading(doc, '2.3.1 界面设计（UI）', 3)
    add_para(doc, '桌面 UI 采用聊天窗口形式，去除了固定快捷按钮，保证所有输入输出都进入 LLM 对话路径。窗口包含历史对话区域、输入框、发送按钮和“思考中”状态提示。这样既符合 Agent 产品形态，也避免把流程固化在界面按钮中。')
    add_picture_with_caption(doc, ui_img, '图2.2 界面原型图')
    add_heading(doc, '2.3.2 流程设计', 3)
    add_para(doc, '系统流程强调观察、思考、行动和验证。LLM 先阅读用户请求和历史上下文，必要时调用 think 记录简短操作意图；随后调用 attach_process、scan_exact_value、next_scan、write_value 等工具。工具结果以结构化 JSON 形式回传给 LLM，模型再决定继续调用工具或向用户回复。')
    add_picture_with_caption(doc, flow_img, '图2.3 系统流程图')

    add_heading(doc, '2.4 系统实现', 2)
    add_heading(doc, '2.4.1 实现效果', 3)
    add_para(doc, '当前系统已经实现 CLI、API 和桌面 UI 三种入口。CLI 支持单次指令和交互模式；API 支持 session_id、多会话状态和 CE 后端占用保护；桌面 UI 提供纯聊天操作界面。系统还提供 check_mcp.py 和 check_llm_tooluse.py 用于检查后端和模型工具调用能力。')
    add_picture_with_caption(doc, effect_img, '图2.5 实现图')

    add_heading(doc, '2.4.2 核心代码', 3)
    add_para(doc, '核心代码一：ToolUseChatAgent 的工具调用循环。该循环负责把用户消息、历史上下文和工具 schema 发送给 LLM，解析模型返回的 tool_calls，并把真实工具执行结果追加回对话。')
    add_code_block(doc, '''for _round in range(self.max_tool_rounds):
    response = self._chat(messages, tools=tool_schemas())
    choice = _normalize_assistant_message(_assistant_message_from_response(response))
    messages.append(choice)
    tool_calls = _normalize_tool_calls(choice)
    if not tool_calls:
        assistant_message = str(choice.get("content") or "我已处理完这轮请求。")
        break
    for tool_call in tool_calls:
        function = tool_call.get("function") or {}
        arguments, parse_error = _parse_tool_arguments(function.get("arguments"))
        action = _action_from_tool_call(str(function.get("name") or ""), arguments)
        result = self.executor.execute(action)
        messages.append({
            "role": "tool",
            "tool_call_id": tool_call.get("id"),
            "content": json.dumps(result_to_tool_payload(result), ensure_ascii=False),
        })''')
    add_para(doc, '核心代码二：CheatEngineMCPExecutor 的写入确认。系统不会只凭写入接口返回就宣布成功，而是继续读回目标地址，并比较读回值与期望值。')
    add_code_block(doc, '''result = self._call("write_integer", {"address": address, "value": value, "type": value_type})
readback = self._call("read_integer", {"address": address, "type": value_type})
write_ok = _write_confirmed(result, readback, value)
if write_ok:
    message = f"已通过 Cheat Engine MCP 将 {label}={value} 写入 {address}。"
else:
    message = f"Cheat Engine MCP 已返回写入结果，但未确认 {label}={value} 已成功写入 {address}。"''')
    add_para(doc, '核心代码三：多会话 CE 后端占用保护。API 入口允许普通聊天并行，但对真实 CE 操作进行所有权检查，避免两个会话同时操作同一个 Cheat Engine MCP 后端。')
    add_code_block(doc, '''def execute(self, action: AgentAction) -> ActionResult:
    if _is_ce_action(action.type) and _ce_session_owner not in (None, self.session_id):
        return ActionResult(
            action=action,
            ok=False,
            message="Cheat Engine MCP backend is currently owned by another API session.",
            data={"error": "ce_session_busy"},
        )
    return self.inner.execute(action)''')

    add_heading(doc, '2.5 系统测试', 2)
    add_test_table(doc)

    add_heading(doc, '2.6 项目总结', 2)
    add_para(doc, 'CheatPilot 项目总体完成了从自然语言任务到真实内存工具执行的完整链路。项目最初的目标是让用户通过对话完成内存数值修改，而不是手动操作传统工具中的扫描、筛选和写入步骤。最终系统实现了 LLM tool-use Agent、Cheat Engine MCP 执行后端、多轮扫描状态管理、CLI/API/UI 三种入口、本地文件和命令工具、模型错误处理以及多项回归测试。通过这些模块，用户可以描述目标程序、当前数值和目标值，系统会根据上下文决定是否需要附加进程、扫描当前值、要求用户改变数值、继续筛选候选地址，并在候选唯一后写入和读回确认。')
    add_para(doc, '项目开发过程中最大的技术难点是把 LLM 的开放式语言能力和真实系统级工具调用结合起来。模型有时会忘记参数、返回旧式 function_call、在工具结果过长时失去上下文，或者因为服务端限流而中断。针对这些问题，我们设计了工具 schema、参数解析错误反馈、工具观察压缩、最大工具轮数限制、历史窗口限制和错误格式化机制。另一个难点是内存扫描状态的连续性。真实扫描往往需要多轮交互，如果用户下一句只说“现在是 120”，系统必须知道这是哪个标签的 next scan。项目通过 session_state 保存标签、候选地址、总匹配数、待写入目标和待打印基址等信息，保证多轮对话不会断链。')
    add_para(doc, '在后端集成方面，项目坚持所有真实内存操作都经过 Cheat Engine MCP，而不是使用 mock 或本地假结果。这样可以验证系统在真实工具链下的可用性，也让错误处理更加重要。我们加入了进程附加后的校验逻辑，避免 Cheat Engine 保留旧进程导致扫描写入错误目标；加入唯一候选判断，避免把单个预览地址误认为总匹配唯一；加入写入读回确认，避免工具返回成功但实际值未变化。API 侧还增加了 CE 后端占用保护，避免多个会话交叉操作。')
    add_para(doc, '不足之处在于，项目目前仍依赖外部 LLM 服务和 Cheat Engine 运行环境，部署和调试门槛较高；指针链分析、复杂结构体识别、可视化日志和更完整的权限隔离仍有扩展空间。后续可以进一步完善模型选择策略，在模型不支持 tool calling 或服务端返回 no_available_providers 时给出更直观的诊断；也可以增加更细粒度的操作日志，方便复盘每一步工具调用。总体来看，本项目使小组成员完整经历了需求分析、技术调研、架构设计、编码实现、真实后端集成、测试验证和文档撰写过程，对 AI Agent 工程化、工具调用协议和系统状态管理有了较深入的理解。')

    doc.save(OUT)
    remove_template_note_shapes(OUT)
    return OUT

if __name__ == '__main__':
    out = build_report()
    print(out)
